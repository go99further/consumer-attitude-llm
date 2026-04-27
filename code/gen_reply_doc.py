from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys
sys.stdout.reconfigure(encoding="utf-8")

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_run(run, bold=False, size=12, color=None, font='宋体'):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 40)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(9)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, bold=True, size=13, color=(0x1F, 0x49, 0x7D), font='黑体')


def add_line(doc, label=None, body=None, bold_parts=None, line_spacing=22, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(line_spacing)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    if label:
        run = p.add_run(label)
        set_run(run, bold=True, size=12, font='黑体')
    if body and not bold_parts:
        run = p.add_run(body)
        set_run(run, size=12)
    if bold_parts:
        for text, bold in bold_parts:
            run = p.add_run(text)
            set_run(run, bold=bold, size=12)
    return p


# ── 主标题 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run('论文修改情况说明')
set_run(run, bold=True, size=18, color=(0x1F, 0x49, 0x7D), font='黑体')

for line in ['致：FYR老师',
             '论文题目：基于在线评论的消费者多属性态度预测效度研究',
             '　　　　　——大语言模型提取方法与评论者经验的调节作用',
             '日期：2026年4月']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(line)
    set_run(run, size=11)

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(22)
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    '感谢老师在上一轮批注中提出的详细意见。您的批注直指论文的核心方法论问题，对我帮助极大。'
    '本次修改以解决CMV问题为核心，对研究设计进行了系统性重构，以下逐条说明修改情况。'
)
set_run(run, size=12)

# ── 批注#2+#10 ──
add_divider(doc)
add_section_heading(doc, '批注#2 + 批注#10　调节变量缺乏理论逻辑')
add_line(doc, label='老师意见：',
         body='用购买行为（verified_purchase）作为调节变量，理论上说不通；建议改用消费者性别、年龄、购物级别等客观数据。')
add_line(doc, label='修改情况：')
add_line(doc, bold_parts=[
    ('已将调节变量从 verified_purchase 更换为', False),
    ('评论者经验（Reviewer Experience）', True),
    ('，以用户在Amazon Beauty品类的历史评论总数衡量。该变量从全品类数据集（23,911,390条记录）中统计每位用户的历史评论数获得，属于平台客观行为数据，与评论文本完全独立。', False),
])
add_line(doc, body=(
    '理论依据为Fazio & Zanna（1978）的直接经验假说和Alba & Hutchinson（1987）的消费者专业性理论。'
    '需要说明的是，Amazon数据集不提供性别、年龄等人口统计信息，评论数量是数据集中可获取的最接近'
    '"消费者专业性"的客观指标，论文中已在局限性章节如实说明这一数据限制。'
))

# ── 批注#3 ──
add_divider(doc)
add_section_heading(doc, '批注#3　未验证购买组发表评论的合理性')
add_line(doc, label='老师意见：', body='未验证购买组为什么可以发表评论？')
add_line(doc, label='修改情况：')
add_line(doc, bold_parts=[
    ('本次修改已', False),
    ('取消已验证/未验证购买的分组对比设计', True),
    ('，将全部1,331条评论统一纳入分析，不再以 verified_purchase 作为分组或调节变量，从根本上规避了这一设计问题。'
     '论文中补充说明了Amazon平台允许通过Vine计划及跨平台购买等机制产生未验证购买评论，作为背景知识说明，不再作为研究变量。', False),
])

# ── 批注#4+#7 ──
add_divider(doc)
add_section_heading(doc, '批注#4 + 批注#7　CMV问题严重，Harman单因子检验78.71%')
add_line(doc, label='老师意见：',
         body='自变量和因变量都来自同一条评论，增加中介变量反而让CMV更严重；真正的解决办法是让至少一个变量来自不同数据源。')
add_line(doc, label='修改情况：')
add_line(doc, bold_parts=[
    ('这是本次修改的核心，采用', False),
    ('三数据源设计', True),
    ('从根源消除CMV：', False),
])

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
headers = ['变量', '数据来源', '是否经过LLM']
rows_data = [
    ['ATT（多属性态度）', '评论文本，LLM提取', '是'],
    ['Rating（评分）', '平台结构化数据', '否'],
    ['Reviewer Experience（评论者经验）', '全品类行为数据', '否'],
]
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
for i, row_data in enumerate(rows_data):
    for j, text in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
add_line(doc, body=(
    '三个核心变量来自三个完全独立的数据源，CMV从根源上得到控制。重新进行Harman单因子检验（三变量），'
    '第一因子解释方差从原来的78.71%降至57.9%，且主要CMV论证依赖方法分离而非Harman数值本身'
    '（Podsakoff et al., 2003）。'
))

# ── 批注#5 ──
add_divider(doc)
add_section_heading(doc, '批注#5　中介变量PV区分度低、因果方向存疑')
add_line(doc, label='老师意见：',
         body='中介变量和其他变量区分度可能非常低；感知价值反而可能影响态度进而影响行为，因果方向存疑。')
add_line(doc, label='修改情况：')
add_line(doc, bold_parts=[
    ('已', False), ('删除PV作为中介变量', True),
    ('。实证数据显示ATT与PV的相关系数r=0.820，两者在语义上高度重叠，难以区分；'
     '且如老师所指出，PV→ATT→BI的因果方向同样合理，无法从理论上确定方向。', False),
])
add_line(doc, body=(
    '论文中明确说明：r=0.820本身即是删除PV的理由——两者测量的是高度重叠的构念，不宜同时纳入模型，'
    '也不宜主张因果关系。PV数据仍保留在数据集中，但不进入任何回归模型。'
))

# ── 批注#6 ──
add_divider(doc)
add_section_heading(doc, '批注#6　区分效度标准太宽松，需要CFA')
add_line(doc, label='老师意见：',
         body='"相关系数<0.85"标准太宽松，需要更稳健的区分效度检验，如验证性因子分析（CFA）。')
add_line(doc, label='修改情况：')
add_line(doc, body=(
    '由于ATT是单一计算值（Σb_i×e_i），PV和BI均为单指标，传统CFA需要多题项量表，不适用于本研究。'
    '因此改用以下组合策略：'
))
for bullet in [
    ('①  预测效度：', 'ATT与Rating（不同方法、不同数据源）的相关r=0.722（p<0.001），证明LLM提取的态度能有效预测客观评分行为。'),
    ('②  区分效度：', 'ATT与reviewer_experience的相关r=0.089（p<0.01），显著低于ATT-Rating相关（r=0.722），参考Fornell & Larcker（1981）和Campbell & Fiske（1959）的效度评估框架。'),
    ('③  增量效度：', 'ATT在控制Rating后仍提供显著增量预测力（ΔR²=0.018，F=8.32，p=0.004），证明多属性分解提供了单一评分之外的额外信息。'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(bullet[0])
    set_run(r1, bold=True, size=12)
    r2 = p.add_run(bullet[1])
    set_run(r2, size=12)

# ── 批注#8 ──
add_divider(doc)
add_section_heading(doc, '批注#8　多重共线性问题')
add_line(doc, label='老师意见：',
         body='ATT、PV高度相关（r=0.820），同时进入回归模型存在严重多重共线性。')
add_line(doc, label='修改情况：')
add_line(doc, bold_parts=[
    ('已删除PV，消除了多重共线性的根源。Study 1回归模型中所有变量', False),
    ('VIF均<1.25', True),
    ('，无多重共线性问题。H2调节效应检验中，对ATT和reviewer_experience进行均值中心化后再创建交互项，VIF同样均<1.25，符合标准。', False),
])

# ── 批注#9 ──
add_divider(doc)
add_section_heading(doc, '批注#9　内部一致性是信度检验，不是效度检验')
add_line(doc, label='老师意见：',
         body="Cronbach's Alpha是信度检验，不能作为效度证据。")
add_line(doc, label='修改情况：')
add_line(doc, body='已纠正概念混淆，在论文中明确区分信度与效度：')
for bullet in [
    ('①  信度（重测信度）：',
     "对50条评论重新独立调用LLM提取，计算两次ATT的Pearson相关系数r=0.892（N=50，p<0.001）；"
     "e_i情感极性一致率97.2%（105/108）。ATT是单一计算值，Cronbach's Alpha不适用，已从论文中删除。"),
    ('②  效度：',
     '预测效度+区分效度+增量效度（见批注#6回复），三者独立于信度检验单独呈现。'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(bullet[0])
    set_run(r1, bold=True, size=12)
    r2 = p.add_run(bullet[1])
    set_run(r2, size=12)

# ── 总结 ──
add_divider(doc)
add_section_heading(doc, '总结')
add_line(doc, body=(
    '本次修改的核心逻辑是：将研究重构为双研究设计——Study 1以三数据源设计从根源消除CMV，'
    '验证LLM方法的预测效度；Study 2保留ATT→BI的理论检验，但诚实报告CMV局限（Harman=76.8%）'
    '和选择偏差（BI存在组ATT显著更高，t=4.60，p<0.001），结论限定于"明确表达行为意向的评论者"子群。'
))
add_line(doc, body='如有不妥之处，恳请老师进一步指正。')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('学生：___________')
set_run(run, size=12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('2026年4月')
set_run(run, size=12)

doc.save(r'E:/bi_ye_she_ji/论文修改回复_致导师.docx')
print('done')
